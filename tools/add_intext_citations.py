from __future__ import annotations

from pathlib import Path
import sys

from docx import Document


CITATIONS = {
    94: "[1][4][5]",
    95: "[1][4]",
    96: "[16][17][18]",
    99: "[1][4]",
    100: "[16][17][18]",
    101: "[15][16][17]",
    102: "[1][2]",
    111: "[16][17]",
    112: "[18]",
    113: "[18][21]",
    118: "[9]",
    121: "[10]",
    124: "[11]",
    132: "[1][5]",
    168: "[9][10][11]",
    194: "[11]",
    240: "[11]",
    275: "[15][16]",
    278: "[1][2]",
    289: "[11]",
    295: "[9][10][11][18]",
    298: "[21][23][24][30]",
}


def inject(text: str, citation: str) -> str:
    if citation in text:
        return text
    stripped = text.rstrip()
    if not stripped:
        return text
    for mark in ("。", ".", "；", ";"):
        if stripped.endswith(mark):
            return stripped[:-1] + citation + mark
    return stripped + citation


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: add_intext_citations.py <source.docx> <target.docx>")
        return 1

    source = Path(sys.argv[1])
    target = Path(sys.argv[2])
    doc = Document(str(source))

    for idx, citation in CITATIONS.items():
        if idx >= len(doc.paragraphs):
            continue
        paragraph = doc.paragraphs[idx]
        paragraph.text = inject(paragraph.text, citation)

    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(target))
    print(target)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
