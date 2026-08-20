from __future__ import annotations

from pathlib import Path
import sys

from docx import Document


def main() -> int:
    if len(sys.argv) != 4:
        print("Usage: dump_paragraph_range.py <docx> <start> <end>")
        return 1

    doc = Document(str(Path(sys.argv[1])))
    start = int(sys.argv[2])
    end = int(sys.argv[3])
    for idx in range(start, min(end + 1, len(doc.paragraphs))):
        p = doc.paragraphs[idx]
        text = p.text.strip().replace("\u3000", " ")
        print(f"{idx}:{p.style.name}:{text}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
