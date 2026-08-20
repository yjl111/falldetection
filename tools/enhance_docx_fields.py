from __future__ import annotations

from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def append_field_run(paragraph, instruction: str, default_text: str = "") -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    r1 = OxmlElement("w:r")
    r1.append(begin)
    paragraph._p.append(r1)

    r2 = OxmlElement("w:r")
    r2.append(instr)
    paragraph._p.append(r2)

    r3 = OxmlElement("w:r")
    r3.append(separate)
    paragraph._p.append(r3)

    if default_text:
        paragraph.add_run(default_text)

    r4 = OxmlElement("w:r")
    r4.append(end)
    paragraph._p.append(r4)


def insert_toc_after_heading(doc: Document) -> None:
    toc_heading = None
    for p in doc.paragraphs:
        if p.text.strip().replace(" ", "") == "目录":
            toc_heading = p
            break

    if toc_heading is None:
        return

    next_p = toc_heading._p.getnext()
    if next_p is not None:
        existing_text = "".join(node.text or "" for node in next_p.iter())
        if "TOC" in existing_text:
            return

    toc_paragraph = doc.add_paragraph()
    toc_heading._p.addnext(toc_paragraph._p)
    append_field_run(toc_paragraph, ' TOC \\\\o "1-3" \\\\h \\\\z \\\\u ', "右键单击以更新目录")


def ensure_page_field(doc: Document) -> None:
    has_page = False
    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            if "PAGE" in paragraph._p.xml:
                has_page = True
                break
        if has_page:
            break

    if has_page:
        return

    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    append_field_run(paragraph, " PAGE ", "1")


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: enhance_docx_fields.py <docx_path>")
        return 1

    path = Path(sys.argv[1])
    doc = Document(str(path))
    insert_toc_after_heading(doc)
    ensure_page_field(doc)
    doc.save(path)
    print(path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
