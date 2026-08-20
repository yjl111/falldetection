from __future__ import annotations

from pathlib import Path
import sys

from docx import Document


def norm(text: str) -> str:
    return text.strip().replace("\u3000", "")


def main() -> int:
    if len(sys.argv) != 2:
        print("Usage: extract_abstract.py <docx>")
        return 1

    doc = Document(str(Path(sys.argv[1])))
    start = None
    end = None

    for idx, para in enumerate(doc.paragraphs):
        text = norm(para.text)
        if start is None and text in {"摘要", "摘要:", "摘 要", "摘要："}:
            start = idx + 1
            continue
        if start is not None and text in {"关键词", "关键词:", "关键词：", "ABSTRACT", "Abstract"}:
            end = idx
            break

    if start is None:
        print("ABSTRACT_NOT_FOUND")
        return 2

    if end is None:
        end = min(start + 10, len(doc.paragraphs))

    for idx in range(start, end):
        text = doc.paragraphs[idx].text.strip()
        if text:
            print(f"{idx}:{text}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
