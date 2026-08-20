from __future__ import annotations

from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor


MARKS = {
    "本文选取跌倒检测系统作为研究对象": "建议为课题意义与应用价值补充文献引用",
    "OpenCV是常用的计算机视觉开源库": "建议为OpenCV功能描述补充文献引用",
    "系统围绕业务流程构建了多组集合": "建议为MongoDB集合设计说明补充文献引用",
    "对于报警留证视频，系统引入MongoDB的GridFS机制进行存储": "建议为GridFS存储机制说明补充文献引用",
    "系统面向两类用户角色提供服务": "建议为认证与权限控制技术说明补充文献引用",
    "为了验证系统的功能完整性与运行稳定性": "建议为测试环境涉及的关键技术栈补充文献引用",
}


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: mark_citation_needed.py <source.docx> <target.docx>")
        return 1

    source = Path(sys.argv[1])
    target = Path(sys.argv[2])
    doc = Document(str(source))

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        for prefix, note in MARKS.items():
            if text.startswith(prefix) and f"【{note}】" not in text:
                for run in paragraph.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                marker = paragraph.add_run(f"【{note}】")
                marker.font.bold = True
                marker.font.color.rgb = RGBColor(192, 0, 0)
                marker.font.highlight_color = WD_COLOR_INDEX.YELLOW
                break

    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(target))
    print(target)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
