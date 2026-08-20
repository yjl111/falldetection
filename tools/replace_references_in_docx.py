from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import sys

from docx import Document


REFERENCES = [
    "[1] 郭晓贝, 王颖, 蒋梦瑶, 等. 基于电子健康的老年人跌倒检测系统的研究进展[J]. 中国临床护理, 2023, 15(10): 649-654. DOI:10.3969/j.issn.1674-3768.2023.10.015.",
    "[2] 徐威, 廖义奎. 基于YOLOv5的空巢老人跌倒检测系统的设计与研究[J]. 软件工程, 2023(2): 31-34.",
    "[3] 周彤彤, 彭月平, 郑璐, 等. 基于YOLOv5s的跌倒行为检测算法[J]. 中国科技论文, 2023, 18(7): 786-792.",
    "[4] 朱云泽. 智慧养老背景下老年人可穿戴设备应用现状分析[J]. 科技创新与应用, 2023, 13(5): 151-155. DOI:10.19981/j.CN23-1581/G3.2023.05.038.",
    "[5] 禹佳景. 城市社区“嵌入式”养老服务模式研究: 以成都市为例[J]. 现代管理, 2023, 13(2): 116-123. DOI:10.12677/MM.2023.132015.",
    "[6] 朱胜豪, 钱承山, 阚希. 改进YOLOv5的高精度跌倒检测算法[J]. 计算机工程与应用, 2024(11): 105-114.",
    "[7] 陈晨, 徐慧英, 朱信忠, 等. 基于YOLOv8改进的室内行人跌倒检测算法FDW-YOLO[J]. 计算机工程与科学, 2024, 46(8): 1455-1465.",
    "[8] 郭夏迪, 曹炳尧. 基于双模态门控特征融合的跌倒检测方法[J]. 计算机测量与控制, 2024, 32(10): 69-76.",
    "[9] 程世通, 张李辉, 楚遵恒, 等. 基于视觉识别和多传感器的跌倒检测系统设计[J]. 传感器与微系统, 2024(8): 91-94.",
    "[10] 张泽滈. 智慧赋能养老服务的驱动要素、转型逻辑、实践困境与对策[J]. 西安交通大学学报(社会科学版), 2024(3): 64-74.",
    "[11] 张语轩, 耿秀丽, 潘飞. 基于多源数据的智慧养老服务供需匹配研究[J]. 上海理工大学学报, 2024(2): 214-224.",
    "[12] 汪睿, 胡珊珊, 陈敏莲, 等. 基于Spring Cloud的分布式医疗数据平台建设[J]. 设备管理与维修, 2024(4): 16-19.",
    "[13] 冯彩英, 张小龙, 李梦婷. 基于多传感器的跌倒检测系统设计与实现[J]. 农业装备与车辆工程, 2025, 63(2): 129-133.",
    "[14] 刘漫, 沈鹏熠, 张茹梦. 人工智能技术在医疗护理中的应用研究[J]. 中国农村卫生事业管理, 2025, 45(3): 177-182.",
    "[15] WANG Y, DENG T. Enhancing elderly care: efficient and reliable real-time fall detection algorithm[J]. Digital Health, 2024, 10: 20552076241233690. DOI:10.1177/20552076241233690.",
    "[16] GAYA-MOREY F X, MANRESA-YEE C, MORA M, et al. Deep learning for computer vision based activity recognition and fall detection of the elderly: a systematic review[J]. Applied Intelligence, 2024, 54: 8982-9007. DOI:10.1007/s10489-024-05645-1.",
    "[17] BENKACI A, SLIMAN L, DELLYS H N. Vision-based human fall detection systems: a review[J]. Procedia Computer Science, 2024, 241: 203-211. DOI:10.1016/j.procs.2024.08.028.",
    "[18] WANG Y, CHI Z, LIU M, et al. High-performance lightweight fall detection with an improved YOLOv5s algorithm[J]. Machines, 2023, 11(8): 818. DOI:10.3390/machines11080818.",
    "[19] ZHENG X, CAO J, WANG C, et al. A high-precision human fall detection model based on FasterNet and deformable convolution[J]. Electronics, 2024, 13(14): 2798. DOI:10.3390/electronics13142798.",
    "[20] QIN Y, MIAO W, QIAN C. A high-precision fall detection model based on dynamic convolution in complex scenes[J]. Electronics, 2024, 13(6): 1141. DOI:10.3390/electronics13061141.",
    "[21] HWANG H, KIM D, KIM H. FD-YOLO: a YOLO network optimized for fall detection[J]. Applied Sciences, 2025, 15(1): 453. DOI:10.3390/app15010453.",
    "[22] SHI H, WANG X, SHI J. Fall detection algorithm using enhanced HRNet combined with YOLO[J]. Sensors, 2025, 25(13): 4128. DOI:10.3390/s25134128.",
    "[23] HUANG X, LI X, YUAN L, et al. SDES-YOLO: a high-precision and lightweight model for fall detection in complex environments[J]. Scientific Reports, 2025, 15: 2026. DOI:10.1038/s41598-025-86593-9.",
    "[24] WANG H, XU S, CHEN Y, et al. LFD-YOLO: a lightweight fall detection network with enhanced feature extraction and fusion[J]. Scientific Reports, 2025, 15: 5069. DOI:10.1038/s41598-025-89214-7.",
    "[25] REN H, LAN P. BMR-YOLO: a deep learning approach for fall detection in complex environments[J]. PLoS One, 2025, 20(11): e0335992. DOI:10.1371/journal.pone.0335992.",
    "[26] PRIYA S, AMSHAKALA K. An adaptive fall detection system based on ensemble learning using variants of YOLO V8, RetinaNet and DETR[J]. Scientific Reports, 2025, 15: 33161. DOI:10.1038/s41598-025-97634-8.",
    "[27] KOTHARI V P, CHAKURKAR P S. Towards safer environments: a YOLO and MediaPipe-based human fall detection system with alert automation[J]. MethodsX, 2025, 15: 103623. DOI:10.1016/j.mex.2025.103623.",
    "[28] JISHNURAJ K, VERGIN RAJA SAROBIN M, ANBARASI J, et al. Fall detection among elderly persons using FallCNN and transfer learning models[J]. Frontiers in Artificial Intelligence, 2026, 9: 1734096. DOI:10.3389/frai.2026.1734096.",
    "[29] ZAFAR R O, ZAFAR F. Real-time activity and fall detection using transformer-based deep learning models for elderly care applications[J]. BMJ Health & Care Informatics, 2025, 32(1): e101439. DOI:10.1136/bmjhci-2025-101439.",
    "[30] AL NAHIAN M J, RAISA J F, MAHMUD M, et al. Artificial intelligence for elderly fall detection: state-of-the-art methods, applications and challenges[J]. Cognitive Computation, 2026, 18: 12. DOI:10.1007/s12559-026-10550-5.",
]


def normalized(text: str) -> str:
    return text.strip().replace(" ", "").replace("\u3000", "")


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: replace_references_in_docx.py <source.docx> <target.docx>")
        return 1

    source = Path(sys.argv[1])
    target = Path(sys.argv[2])
    doc = Document(str(source))
    paragraphs = doc.paragraphs

    ref_heading_idx = None
    end_idx = None
    for idx, paragraph in enumerate(paragraphs):
        text = normalized(paragraph.text)
        if ref_heading_idx is None and text == "参考文献":
            ref_heading_idx = idx
            continue
        if ref_heading_idx is not None and text in {"致谢", "鸣谢", "附录", "附录A", "附录B"}:
            end_idx = idx
            break

    if ref_heading_idx is None or end_idx is None:
        print("REFERENCE_SECTION_NOT_FOUND")
        return 2

    template_paragraph = paragraphs[ref_heading_idx + 1]
    end_paragraph = paragraphs[end_idx]

    for idx in range(end_idx - 1, ref_heading_idx, -1):
        p = paragraphs[idx]._element
        p.getparent().remove(p)

    for ref in REFERENCES:
        new_p = deepcopy(template_paragraph._element)
        for child in list(new_p):
            # Keep paragraph properties, drop existing runs
            if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
                new_p.remove(child)
        end_paragraph._element.addprevious(new_p)
        inserted = doc.paragraphs[0]
        for p in doc.paragraphs:
            if p._element is new_p:
                inserted = p
                break
        inserted.text = ref

    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(target))
    print(target)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
