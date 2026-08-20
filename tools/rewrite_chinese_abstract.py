from __future__ import annotations

from pathlib import Path
import sys

from docx import Document


NEW_ABSTRACT = [
    "随着人口老龄化加深，老年人居家与机构养老场景中的安全监护需求不断提升。跌倒事件具有突发性强、危害性高和救援时效要求高等特点，若不能及时发现，容易造成严重后果。相比依赖主动佩戴的传统穿戴式监测方案，基于计算机视觉的非接触式跌倒检测方式更适合连续监护场景。",
    "本文面向智能养老应用需求，设计并实现了一套基于深度学习的跌倒检测与预警系统。系统采用前后端分离架构，前端基于 Vue 3 实现可视化交互，后端基于 Flask 提供业务接口与视频处理服务，核心检测环节结合 YOLO 模型与 OpenCV 对摄像头或本地视频中的目标行为进行实时识别。",
    "在业务实现上，系统完成了实时检测、报警管理、历史回放、系统设置和模型训练等基础功能，并扩展了角色权限控制、用户资料、报警工单、通知日志、健康报告、设备心跳和通知规则等模块。数据层采用 MongoDB 存储业务数据，并结合 GridFS 保存报警留证视频，形成了覆盖报警记录、视频片段、关键帧截图、用户消息和审计日志的业务闭环。研究结果表明，该系统能够满足跌倒检测场景下实时监测、信息留存和事件处理的基本需求，可为家庭养老和社区看护场景提供参考。",
]


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: rewrite_chinese_abstract.py <source.docx> <target.docx>")
        return 1

    doc = Document(str(Path(sys.argv[1])))

    for idx, text in zip((15, 16, 17), NEW_ABSTRACT):
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].text = text

    target = Path(sys.argv[2])
    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(target))
    print(target)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
